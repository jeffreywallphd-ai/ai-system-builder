from __future__ import annotations

import html
from dataclasses import dataclass
from pathlib import Path
import re
from zipfile import BadZipFile, ZipFile

from ..models import (
    DatasetPreparationSourceInput,
    DatasetPreparationWarning,
    DocumentNormalizationConfig,
)


@dataclass
class NormalizedRegion:
    kind: str
    start: int
    end: int
    page_number: int | None = None
    confidence: float = 1.0


@dataclass
class NormalizedDocument:
    artifact_id: str
    markdown: str
    media_type: str | None
    source_path: str
    regions: list[NormalizedRegion] | None = None
    extraction_quality: float = 1.0


@dataclass
class _NormalizedContent:
    markdown: str
    regions: list[NormalizedRegion]
    extraction_quality: float


@dataclass
class DocumentNormalizationResult:
    documents: list[NormalizedDocument]
    skipped_document_count: int
    warnings: list[DatasetPreparationWarning]


_SUPPORTED_SUFFIXES = {
    ".txt",
    ".md",
    ".markdown",
    ".html",
    ".htm",
    ".pdf",
    ".docx",
    ".csv",
    ".json",
    ".jsonl",
}

_UNSUPPORTED_BUT_COMMON_SUFFIXES = {".doc"}
MAX_DOCUMENT_SOURCE_BYTES = 256 * 1024 * 1024
MAX_EXTRACTED_DOCUMENT_CHARACTERS = 20_000_000
MAX_PDF_PAGE_COUNT = 5_000
MAX_DOCX_ARCHIVE_ENTRY_COUNT = 10_000
MAX_DOCX_UNCOMPRESSED_BYTES = 512 * 1024 * 1024


def _extension_for_source(source: DatasetPreparationSourceInput) -> str:
    if source.originalName:
        original_extension = Path(source.originalName).suffix.lower()
        if original_extension:
            return original_extension

    return Path(source.localPath).suffix.lower()


def _read_text(path: Path) -> str:
    text = path.read_text(encoding="utf-8")
    _assert_extracted_text_bound(text)
    return text


def _assert_extracted_text_bound(text: str) -> None:
    if len(text) > MAX_EXTRACTED_DOCUMENT_CHARACTERS:
        raise ValueError("The document contains more extracted text than can be prepared safely.")


def _extraction_quality(text: str) -> float:
    if not text.strip():
        return 0.0
    replacement_ratio = text.count("\ufffd") / max(1, len(text))
    control_count = sum(1 for character in text if ord(character) < 32 and character not in "\n\r\t")
    control_ratio = control_count / max(1, len(text))
    return round(max(0.0, 1.0 - min(1.0, replacement_ratio * 10 + control_ratio * 10)), 4)


def _regions_for_text(text: str, *, page_number: int | None = None) -> list[NormalizedRegion]:
    regions: list[NormalizedRegion] = []
    for match in re.finditer(r"[^\n]+", text):
        value = match.group(0)
        stripped = value.strip()
        if not stripped:
            continue
        leading = len(value) - len(value.lstrip())
        start = match.start() + leading
        end = start + len(stripped)
        kind = "paragraph"
        if stripped.startswith("#"):
            kind = "heading"
        elif stripped.count("|") >= 2:
            kind = "table"
        regions.append(
            NormalizedRegion(
                kind=kind,
                start=start,
                end=end,
                page_number=page_number,
            )
        )
    if not regions and text:
        regions.append(NormalizedRegion(kind="text", start=0, end=len(text), page_number=page_number))
    return regions


def _content_from_text(text: str) -> _NormalizedContent:
    _assert_extracted_text_bound(text)
    return _NormalizedContent(
        markdown=text,
        regions=_regions_for_text(text),
        extraction_quality=_extraction_quality(text),
    )


def _normalize_html(path: Path) -> _NormalizedContent:
    try:
        from markdownify import markdownify as html_to_markdown
    except ImportError as error:  # pragma: no cover - covered through policy behavior
        raise RuntimeError("markdownify is required for HTML normalization") from error

    return _content_from_text(html_to_markdown(_read_text(path), heading_style="ATX"))


def _normalize_pdf(path: Path) -> _NormalizedContent:
    try:
        from pypdf import PdfReader
    except ImportError as error:  # pragma: no cover - covered through policy behavior
        raise RuntimeError("pypdf is required for PDF normalization") from error

    reader = PdfReader(str(path))
    if len(reader.pages) > MAX_PDF_PAGE_COUNT:
        raise ValueError("The PDF contains too many pages to prepare safely.")
    page_text: list[str] = []
    regions: list[NormalizedRegion] = []
    extracted_character_count = 0
    for page_index, page in enumerate(reader.pages):
        extracted = page.extract_text() or ""
        if extracted.strip():
            normalized = extracted.strip()
            extracted_character_count += len(normalized)
            if extracted_character_count > MAX_EXTRACTED_DOCUMENT_CHARACTERS:
                raise ValueError("The PDF contains more extracted text than can be prepared safely.")
            start = sum(len(value) + 2 for value in page_text)
            page_text.append(normalized)
            end = start + len(normalized)
            regions.append(
                NormalizedRegion(
                    kind="page",
                    start=start,
                    end=end,
                    page_number=page_index + 1,
                    confidence=0.9,
                )
            )

    text = "\n\n".join(page_text)
    return _NormalizedContent(
        markdown=text,
        regions=regions,
        extraction_quality=_extraction_quality(text),
    )


def _normalize_docx(path: Path) -> _NormalizedContent:
    try:
        from docx import Document
    except ImportError as error:  # pragma: no cover - covered through policy behavior
        raise RuntimeError("python-docx is required for DOCX normalization") from error

    try:
        with ZipFile(path) as archive:
            entries = archive.infolist()
            if len(entries) > MAX_DOCX_ARCHIVE_ENTRY_COUNT:
                raise ValueError("The Word document contains too many embedded files.")
            if sum(entry.file_size for entry in entries) > MAX_DOCX_UNCOMPRESSED_BYTES:
                raise ValueError("The Word document expands beyond the safe preparation limit.")
    except BadZipFile as error:
        raise ValueError("The Word document is not a valid DOCX file.") from error

    document = Document(str(path))
    lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        table_rows = [
            "| " + " | ".join(cell.text.strip() for cell in row.cells) + " |"
            for row in table.rows
        ]
        if table_rows:
            lines.extend(table_rows)
    normalized = "\n\n".join(lines)
    return _content_from_text(normalized)


def _normalize_source_to_markdown(source: DatasetPreparationSourceInput) -> _NormalizedContent:
    path = Path(source.localPath)
    if not path.exists():
        raise ValueError("The staged input file is not available.")
    if not path.is_file():
        raise ValueError("The staged input is not a file.")
    if path.stat().st_size > MAX_DOCUMENT_SOURCE_BYTES:
        raise ValueError("The document exceeds the safe preparation size limit.")

    extension = _extension_for_source(source)
    if extension in _UNSUPPORTED_BUT_COMMON_SUFFIXES:
        raise ValueError(
            "Unsupported document type: .doc (legacy Microsoft Word). Convert to .docx before dataset preparation."
        )

    if extension not in _SUPPORTED_SUFFIXES:
        raise ValueError(f"Unsupported document type: {extension or 'unknown'}")

    if extension in {".md", ".markdown"}:
        return _content_from_text(_read_text(path))

    if extension in {".txt", ".csv", ".json", ".jsonl"}:
        return _content_from_text(_read_text(path))

    if extension in {".html", ".htm"}:
        return _normalize_html(path)

    if extension == ".pdf":
        return _normalize_pdf(path)

    if extension == ".docx":
        return _normalize_docx(path)

    raise ValueError(f"Unsupported document type: {html.escape(extension)}")


def normalize_sources_to_markdown(
    source_inputs: list[DatasetPreparationSourceInput],
    config: DocumentNormalizationConfig,
) -> DocumentNormalizationResult:
    if config.targetFormat != "markdown":
        raise ValueError(f"Unsupported normalization target format: {config.targetFormat}")

    policy = config.unsupportedDocumentPolicy or "fail"
    warnings: list[DatasetPreparationWarning] = []
    normalized: list[NormalizedDocument] = []
    skipped = 0

    for source in source_inputs:
        try:
            content = _normalize_source_to_markdown(source)
            normalized.append(
                NormalizedDocument(
                    artifact_id=source.artifactId,
                    markdown=content.markdown,
                    media_type=source.mediaType,
                    source_path=source.localPath,
                    regions=content.regions,
                    extraction_quality=content.extraction_quality,
                )
            )
        except Exception as error:
            if policy == "skip":
                skipped += 1
                warning_code = "document_normalization_skipped"
                if ".doc (legacy Microsoft Word)" in str(error):
                    warning_code = "document_normalization_unsupported_doc"
                warnings.append(
                    DatasetPreparationWarning(
                        code=warning_code,
                        message=(
                            f"Skipped source '{source.artifactId}' because it could not be read "
                            "with the selected document settings."
                        ),
                        sourceArtifactId=source.artifactId,
                    )
                )
                continue
            raise

    return DocumentNormalizationResult(
        documents=normalized,
        skipped_document_count=skipped,
        warnings=warnings,
    )
